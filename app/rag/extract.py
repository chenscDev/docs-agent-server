"""从本地文件提取纯文本。"""

from __future__ import annotations

from collections.abc import Callable
from pathlib import Path

from app.core.config import get_settings

# 允许的扩展名（小写）
ALLOWED_EXTENSIONS = {".pdf", ".txt", ".md", ".markdown", ".docx"}

MIME_BY_EXT = {
    ".pdf": "application/pdf",
    ".txt": "text/plain",
    ".md": "text/markdown",
    ".markdown": "text/markdown",
    ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
}

# 演示/练手期 PDF 页数上限（过大易拖垮本机 embedding）
MAX_PDF_PAGES = 100

StageCallback = Callable[[str, float | None], None]


class ExtractError(Exception):
    """文本提取失败。"""

    def __init__(self, code: str, message: str) -> None:
        self.code = code
        self.message = message
        super().__init__(message)


def normalize_extension(filename: str) -> str:
    """返回小写扩展名，含点号。"""
    return Path(filename).suffix.lower()


def guess_mime(filename: str) -> str:
    """根据扩展名推断 MIME。"""
    ext = normalize_extension(filename)
    return MIME_BY_EXT.get(ext, "application/octet-stream")


def extract_text(
    file_path: Path,
    filename: str,
    *,
    on_stage: StageCallback | None = None,
) -> str:
    """
    提取文本。

    成功返回非空字符串；失败抛 ExtractError（带稳定 code）。
    on_stage：可选进度回调 (message, progress|None)，供 OCR 页进度刷新。
    """
    ext = normalize_extension(filename)
    if ext not in ALLOWED_EXTENSIONS:
        raise ExtractError("UNSUPPORTED_TYPE", f"不支持的文件类型: {ext or '(无扩展名)'}")

    if not file_path.is_file():
        raise ExtractError("INTERNAL", f"文件不存在: {file_path}")

    try:
        if ext == ".pdf":
            text = _extract_pdf(file_path, on_stage=on_stage)
        elif ext == ".docx":
            text = _extract_docx(file_path)
        else:
            text = _extract_plain_text(file_path)
    except ExtractError:
        raise
    except Exception as exc:  # noqa: BLE001
        raise ExtractError("INTERNAL", f"提取文本失败: {exc}") from exc

    cleaned = text.strip()
    if not cleaned:
        raise ExtractError(
            "PARSE_EMPTY",
            "未能提取到可用文本（可能是扫描版 PDF、空文件或仅图片页）",
        )
    return cleaned


def _effective_char_count(text: str) -> int:
    """统计有效字符（去空白），用于判断文本层是否「够用」。"""
    return len("".join((text or "").split()))


def _extract_plain_text(file_path: Path) -> str:
    """读取 txt/md；拒绝明显二进制。"""
    raw = file_path.read_bytes()
    if b"\x00" in raw[:4096]:
        raise ExtractError("PARSE_BINARY", "文本文件疑似二进制，无法按 UTF-8 解析")
    return raw.decode("utf-8", errors="ignore")


def _extract_docx(file_path: Path) -> str:
    """抽取 docx 段落与表格文本。"""
    try:
        from docx import Document as DocxDocument
    except ModuleNotFoundError as exc:
        raise ExtractError(
            "PARSE_DEPENDENCY",
            "缺少 python-docx。请执行: .venv/bin/pip install python-docx",
        ) from exc

    doc = DocxDocument(str(file_path))
    parts: list[str] = []
    for para in doc.paragraphs:
        line = (para.text or "").strip()
        if line:
            parts.append(line)

    for table in doc.tables:
        for row in table.rows:
            cells = [(c.text or "").strip() for c in row.cells]
            # 去空单元格后用制表符拼接，保留表结构信息
            line = "\t".join(c for c in cells if c)
            if line:
                parts.append(line)

    return "\n".join(parts)


def _extract_pdf(
    file_path: Path,
    *,
    on_stage: StageCallback | None = None,
) -> str:
    """使用 PyMuPDF 提取 PDF 文本；不足时走云 OCR。"""
    try:
        import fitz  # pymupdf
    except ModuleNotFoundError as exc:
        raise ExtractError(
            "PARSE_DEPENDENCY",
            "缺少 PyMuPDF（import fitz 失败）。请用仓库 .venv 启动服务："
            ".venv/bin/pip install -r requirements.txt",
        ) from exc

    settings = get_settings()
    min_chars = max(1, int(settings.ocr_min_chars))

    with fitz.open(file_path) as doc:
        if doc.is_encrypted:
            # 尝试空密码（部分 PDF 仅权限加密）
            auth_ok = False
            try:
                auth_ok = bool(doc.authenticate(""))
            except Exception:  # noqa: BLE001
                auth_ok = False
            if not auth_ok:
                raise ExtractError("PDF_ENCRYPTED", "PDF 已加密，请提供未加密文件")

        page_count = int(doc.page_count or 0)
        if page_count <= 0:
            raise ExtractError("PARSE_EMPTY", "PDF 无页面")

        if page_count > MAX_PDF_PAGES:
            raise ExtractError(
                "PDF_TOO_MANY_PAGES",
                f"PDF 页数 {page_count} 超过上限 {MAX_PDF_PAGES}，请拆分后再上传",
            )

        parts: list[str] = []
        for page in doc:
            parts.append(page.get_text("text") or "")

    layer_text = "\n".join(parts)
    if _effective_char_count(layer_text) >= min_chars:
        return layer_text

    # 文本层不足：尝试 OCR
    if not settings.ocr_enabled:
        raise ExtractError(
            "PARSE_EMPTY",
            "未能提取到可用文本（可能是扫描版 PDF）。当前已关闭 OCR（OCR_ENABLED=false）",
        )

    from app.rag.ocr import OcrError, ocr_pdf_file

    if on_stage is not None:
        on_stage("文本层不足，开始 OCR…", 0.15)

    try:
        return ocr_pdf_file(file_path, on_progress=on_stage)
    except OcrError as exc:
        raise ExtractError(exc.code, exc.message) from exc
